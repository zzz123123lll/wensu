"""统一导出服务（P0-5）：普通导出与 Review 导出共享同一套数据装配。

- 装配：标题 + 正文 blocks + 引用清单（来源标题/URL/证据原文/核验状态）+ 来源附录
- 格式：Markdown / 纯文本 / Word（合法 DOCX，python-docx）
- 诚实显示：orphaned / source_dead / needs_recheck 等状态不伪装为已支持
- 导出只读：不修改原稿、不落库（Review 导出另存 manifest 由 review 模块负责）
- 文件名安全：拒绝路径穿越与非法字符
"""

import io
import json
import re
import zipfile
from datetime import datetime, timezone

from app import blocks as blocks_lib
from app import db

# 核验状态 → 用户可读标签（与 db.VERIF_* 一致；诚实展示，不美化）
VERIF_LABELS = {
    "pending": "待核验",
    "supported": "已核验",
    "insufficient": "证据不足",
    "conflicting": "证据冲突",
    "source_dead": "来源失效",
    "needs_recheck": "需复查（正文已变化）",
}
CITATION_STATUS_LABELS = {
    "active": "",
    "orphaned": "（正文已删除，孤立引用）",
}


class ExportError(Exception):
    pass


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _verif_label(c) -> str:
    v = c.get("verification_status") or "pending"
    label = VERIF_LABELS.get(v, f"未知状态({v})")
    status_label = CITATION_STATUS_LABELS.get(c.get("status"), "")
    return (label + status_label) if status_label else label


def build_export_data(conn, aid: int) -> dict:
    """统一装配：文章 + 引用数据 + 编号映射。

    - 引用编号 = citations 列表顺序（1 起，与附录顺序一致）
    - 每个引用含：来源标题、URL、证据原文片段、核验状态、引用状态
    """
    art = db.get_article(conn, aid)
    if art is None:
        raise ExportError(f"草稿 {aid} 不存在")
    citations = db.list_citations(conn, aid)
    numbered = [(i + 1, c) for i, c in enumerate(citations)]
    by_block: dict[str, list[tuple[int, dict]]] = {}
    for n, c in numbered:
        bid = c.get("block_id")
        if bid:
            by_block.setdefault(bid, []).append((n, c))
    return {
        "article": art,
        "citations": numbered,
        "by_block": by_block,
    }


def _block_markdown(b: dict) -> str:
    try:
        return blocks_lib.serialize_blocks([b]).rstrip("\n")
    except Exception:
        return b.get("text", "")


def _citation_entry_md(n: int, c: dict) -> str:
    title = c.get("source_title") or c.get("source_url") or "（无标题来源）"
    url = c.get("source_url") or ""
    quote = (c.get("quote") or "").strip()
    verif = _verif_label(c)
    parts = [f"[{n}] {title}"]
    if url:
        parts.append(url)
    if quote:
        parts.append(f"证据：{quote[:200]}")
    parts.append(f"核验：{verif}")
    return "；".join(parts)


def render_markdown(data: dict, include_appendix: bool = True) -> str:
    """Markdown：标题 + 正文（引用编号上标）+ 引用清单 + 可选来源附录。"""
    art = data["article"]
    citations = data["citations"]
    lines = [f"# {art['title']}", ""]
    for b in art["blocks"]:
        md = _block_markdown(b)
        nums = data["by_block"].get(b.get("id"), [])
        if nums:
            md += " <sup>" + "".join(f"[{n}]" for n, _ in nums) + "</sup>"
        if md:
            lines.append(md)
            lines.append("")
    if citations:
        lines.append("---")
        lines.append("")
        lines.append("## 引用清单")
        lines.append("")
        for n, c in citations:
            lines.append(_citation_entry_md(n, c))
        if include_appendix:
            lines.append("")
            lines.append("---")
            lines.append("")
            lines.append("## 来源附录")
            lines.append("")
            for n, c in citations:
                url = c.get("source_url") or ""
                lines.append(f"- [{n}] {c.get('source_title') or url or '（无标题）'}"
                             + (f"：{url}" if url else ""))
    return "\n".join(lines).rstrip() + "\n"


def render_plain(data: dict, include_appendix: bool = True) -> str:
    """纯文本：标题 + 正文（引用编号 [N] 行内）+ 引用清单 + 可选来源附录。"""
    art = data["article"]
    citations = data["citations"]
    lines = [art["title"], ""]
    for b in art["blocks"]:
        text = b.get("text", "")
        nums = data["by_block"].get(b.get("id"), [])
        if nums:
            text += " [" + "][".join(str(n) for n, _ in nums) + "]"
        if text:
            lines.append(text)
            lines.append("")
    if citations:
        lines.append("引用清单")
        lines.append("=" * 8)
        lines.append("")
        for n, c in citations:
            lines.append(_citation_entry_md(n, c))
            lines.append("")
        if include_appendix:
            lines.append("来源附录")
            lines.append("=" * 8)
            lines.append("")
            for n, c in citations:
                url = c.get("source_url") or ""
                lines.append(f"[{n}] {c.get('source_title') or url or '（无标题）'}"
                             + (f"：{url}" if url else ""))
    return "\n".join(lines).rstrip() + "\n"


def render_docx(data: dict, include_appendix: bool = True) -> bytes:
    """Word：合法 DOCX（python-docx 生成，非改扩展名）。"""
    from docx import Document
    from docx.shared import Pt

    art = data["article"]
    citations = data["citations"]
    doc = Document()
    doc.add_heading(art["title"], level=0)
    for b in art["blocks"]:
        btype = b.get("type", "paragraph")
        text = b.get("text", "")
        nums = data["by_block"].get(b.get("id"), [])
        if nums:
            text += " [" + "][".join(str(n) for n, _ in nums) + "]"
        if btype in ("heading", "heading2", "heading3", "heading4"):
            level = {"heading": 1, "heading2": 2, "heading3": 3, "heading4": 4}[btype]
            doc.add_heading(text, level=level)
        elif btype == "blockquote":
            p = doc.add_paragraph(text, style="Intense Quote")
        elif btype in ("unordered_list", "ordered_list"):
            for item in text.split("\n"):
                doc.add_paragraph(item, style="List Bullet" if btype == "unordered_list" else "List Number")
        elif btype == "code":
            p = doc.add_paragraph(text)
            p.style = doc.styles["No Spacing"]
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        else:
            doc.add_paragraph(text)
    if citations:
        doc.add_heading("引用清单", level=1)
        for n, c in citations:
            doc.add_paragraph(_citation_entry_md(n, c))
        if include_appendix:
            doc.add_heading("来源附录", level=1)
            for n, c in citations:
                url = c.get("source_url") or ""
                doc.add_paragraph(f"[{n}] {c.get('source_title') or url or '（无标题）'}"
                                  + (f"：{url}" if url else ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _citation_entry_wechat(n: int, c: dict) -> str:
    """引用条目（Markdown 形态，供公众号转换器渲染为可点击链接）。"""
    title = c.get("source_title") or c.get("source_url") or "（无标题来源）"
    url = c.get("source_url") or ""
    label = f"[{n}] {title}"
    if url:
        label = f"[{n}] [{title}]({url})"
    return f"- {label}（核验：{_verif_label(c)}）"


def render_wechat(data: dict, theme: str = "default", include_appendix: bool = True) -> str:
    """公众号 HTML：正文（行内样式片段）+ 引用清单 + 可选来源附录。

    - 引用标记用纯文本 [N]（微信编辑器对 <sup> 支持差）
    - 标题不进入正文（公众号标题填在标题栏，避免重复）
    """
    from app.domains import wechat_html

    art = data["article"]
    lines = []
    for b in art["blocks"]:
        md = _block_markdown(b)
        nums = data["by_block"].get(b.get("id"), [])
        if nums:
            md += " " + "".join(f"[{n}]" for n, _ in nums)
        if md:
            lines.append(md)
            lines.append("")
    body = wechat_html.markdown_to_wechat_html("\n".join(lines), theme=theme)
    out = [body]
    if data["citations"]:
        tail = ["---", "", "## 引用清单", ""]
        tail.extend(_citation_entry_wechat(n, c) for n, c in data["citations"])
        if include_appendix:
            tail += ["", "---", "", "## 来源附录", ""]
            for n, c in data["citations"]:
                url = c.get("source_url") or ""
                title = c.get("source_title") or url or "（无标题）"
                tail.append(f"- [{n}] [{title}]({url})" if url else f"- [{n}] {title}")
        out.append(wechat_html.markdown_to_wechat_html("\n".join(tail), theme=theme))
    return "\n".join(part for part in out if part) + "\n"


RENDERERS = {
    "md": render_markdown,
    "markdown": render_markdown,
    "txt": render_plain,
    "text": render_plain,
    "docx": render_docx,
    "wechat": render_wechat,
}


def render(data: dict, fmt: str, include_appendix: bool = True, theme: str = "default") -> bytes:
    renderer = RENDERERS.get(fmt.lower())
    if renderer is None:
        raise ExportError(f"不支持的导出格式: {fmt}")
    if fmt.lower() == "wechat":
        out = renderer(data, theme=theme, include_appendix=include_appendix)
    else:
        out = renderer(data, include_appendix)
    return out if isinstance(out, bytes) else out.encode("utf-8")


# ---------- 文件名安全 ----------

_UNSAFE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def safe_filename(title: str, fmt: str, existing: list[str] | None = None) -> str:
    """安全文件名：清理非法字符、连续点号、拒绝路径穿越、冲突追加时间戳。"""
    cleaned = _UNSAFE.sub("_", (title or "文章").strip()) or "文章"
    cleaned = re.sub(r"\.{2,}", ".", cleaned)  # 连续点号（如 a..b）在部分文件系统有特殊语义
    cleaned = cleaned[:40]
    ext = "docx" if fmt == "docx" else ("txt" if fmt == "txt" else ("html" if fmt in ("wechat",) else ("zip" if fmt == "zip" else "md")))
    name = f"{cleaned}-{_now()[:10]}.{ext}"
    if existing and name in existing:
        ts = datetime.now().strftime("%H%M%S")
        name = f"{cleaned}-{_now()[:10]}-{ts}.{ext}"
    if ".." in name or name.startswith("/") or name.startswith("\\"):
        raise ValueError("非法文件名")
    return name


# ---------- 项目级导出（P2-⑩） ----------

def project_name(conn, pid: int) -> str:
    """项目名（不存在抛 ExportError）。"""
    for p in db.list_projects(conn):
        if p[0] == pid:
            return p[1]
    raise ExportError(f"项目 {pid} 不存在")


def build_project_export(conn, pid: int) -> bytes:
    """项目级 ZIP：manifest + 每篇 Markdown + 素材清单 + 来源清单。只读导出。"""
    name = project_name(conn, pid)
    arts = db.list_articles(conn, pid)
    materials = db.list_materials(conn, project_id=pid)
    sources = db.list_sources(conn, pid)
    root = name[:30] or "项目"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        manifest = {
            "product": "文序",
            "project": name,
            "exported_at": _now(),
            "articles": [],
            "materials_count": len(materials),
            "sources_count": len(sources),
        }
        for aid, title, _updated in arts:
            data = build_export_data(conn, aid)
            md = render_markdown(data)
            zf.writestr(f"{root}/文章/{safe_filename(title, 'md')}", md)
            manifest["articles"].append({
                "id": aid,
                "title": title,
                "version": data["article"]["version"],
                "citations_count": len(data["citations"]),
            })
        if materials:
            lines = ["# 素材清单", ""]
            for m in materials:
                lines.append(f"## {m['title']}")
                lines.append("")
                lines.append(str(m.get("content") or "")[:2000])
                lines.append("")
                tags = m.get("tags") or []
                if tags:
                    lines.append("标签：" + "、".join(str(t) for t in tags))
                    lines.append("")
            zf.writestr(f"{root}/素材清单.md", "\n".join(lines))
        if sources:
            lines = ["# 来源清单", ""]
            for s in sources:
                title = s.get("title") or "（无标题）"
                url = s.get("url") or s.get("canonical_url") or ""
                lines.append(f"- {title}" + (f"：{url}" if url else ""))
            zf.writestr(f"{root}/来源清单.md", "\n".join(lines))
        zf.writestr(f"{root}/manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2))
    return buf.getvalue()
